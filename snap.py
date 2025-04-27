import tkinter as tk
from tkinter import messagebox, filedialog
from openpyxl import Workbook, load_workbook
from datetime import datetime
from docx import Document
from pathlib import Path
import os

class MSnapApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("M-SNAP Form")
        self.geometry("600x650")
        self.resizable(False, False)

        self.dropdown_fields_common = {
            "Stray?": ["Yes", "No", "N/A"],
            "Older than 5 years?": ["Yes", "No", "N/A"]
        }
        self.gender_field = {"Gender": ["Male", "Female", "N/A"]}
        self.yes_no_field = {"Within Morg city limits?": ["Yes", "No"]}

        self.init_data()
        self.build_frames()
        self.show_frame(0)

    def init_data(self):
        self.frames = []
        self.current_frame = 0
        self.add_another_pet_flags = {}

        self.data = {
            "Recipient Information": {
                "First Name": tk.StringVar(),
                "Last Name": tk.StringVar(),
                "Day phone(s)": tk.StringVar(),
                "Street, Apt # OR PO Box": tk.StringVar(),
                "City": tk.StringVar(),
                "Zip": tk.StringVar(),
                "Within Morg city limits?": tk.StringVar(value="No"),
                "POB only: closest town": tk.StringVar(),
                "How did you hear about M-SNAP?": tk.StringVar(),
                "Name on voucher": tk.StringVar()
            }
        }
        for i in range(1, 4):
            section = f"Pet {i} Information"
            self.data[section] = {
                "Name": tk.StringVar(),
                "Species": tk.StringVar(),
                "Gender": tk.StringVar(value="N/A"),
                "Breed": tk.StringVar(),
                "Color(s)": tk.StringVar(),
                "Distinguishing characteristic(s)": tk.StringVar(),
                "Stray?": tk.StringVar(value="N/A"),
                "Older than 5 years?": tk.StringVar(value="N/A"),
                "Dog weight range?": tk.StringVar(),  # Now text input
                "Special": tk.StringVar(),            # Now text input
                "Voucher": tk.StringVar(),
                "Sent": tk.StringVar(),
                "Expires": tk.StringVar(),
                "Grant": tk.StringVar()
            }

    def build_frames(self):
        for idx, (section, fields) in enumerate(self.data.items()):
            frame = tk.Frame(self, padx=20, pady=20)
            tk.Label(frame, text=section, font=('Arial', 16, 'bold')).grid(row=0, column=0, columnspan=2, pady=(0,20))
            row = 1
            for label, var in fields.items():
                tk.Label(frame, text=label, width=30, anchor='w').grid(row=row, column=0, padx=5, pady=5, sticky='w')
                if section == "Recipient Information" and label in self.yes_no_field:
                    opt = tk.OptionMenu(frame, var, *self.yes_no_field[label])
                    opt.config(width=38)
                    opt.grid(row=row, column=1, padx=5, pady=5)
                elif section.startswith("Pet") and label in self.dropdown_fields_common:
                    opt = tk.OptionMenu(frame, var, *self.dropdown_fields_common[label])
                    opt.config(width=38)
                    opt.grid(row=row, column=1, padx=5, pady=5)
                elif section.startswith("Pet") and label in self.gender_field:
                    opt = tk.OptionMenu(frame, var, *self.gender_field[label])
                    opt.config(width=38)
                    opt.grid(row=row, column=1, padx=5, pady=5)
                else:
                    tk.Entry(frame, textvariable=var, width=40).grid(row=row, column=1, padx=5, pady=5)
                row += 1

            if section in ["Pet 1 Information", "Pet 2 Information"]:
                flag = tk.IntVar()
                self.add_another_pet_flags[section] = flag
                tk.Checkbutton(frame, text="Add another pet?", variable=flag).grid(row=row, column=0, columnspan=2, pady=(10,5))
                row += 1

            btn_frame = tk.Frame(frame)
            btn_frame.grid(row=row, column=0, columnspan=2, pady=20)
            if idx > 0:
                tk.Button(btn_frame, text="Back", command=self.prev_frame).pack(side=tk.LEFT, padx=5)
            if section in ["Pet 1 Information", "Pet 2 Information"]:
                tk.Button(btn_frame, text="Next", command=lambda sec=section: self.handle_pet_next(sec)).pack(side=tk.LEFT, padx=5)
            elif section == "Recipient Information":
                tk.Button(btn_frame, text="Next", command=self.next_frame).pack(side=tk.LEFT, padx=5)
            else:
                tk.Button(btn_frame, text="Submit", command=self.export_to_excel).pack(side=tk.LEFT, padx=5)

            self.frames.append(frame)

    def handle_pet_next(self, section):
        if self.add_another_pet_flags[section].get():
            nxt = 2 if section.endswith("1 Information") else 3
            self.show_frame(nxt)
        else:
            self.choose_output_folder()


    def show_frame(self, idx):
        for f in self.frames:
            f.grid_forget()
        self.frames[idx].grid(row=0, column=0, sticky='nsew')
        self.current_frame = idx

    def next_frame(self):
        if self.current_frame < len(self.frames)-1:
            self.show_frame(self.current_frame+1)

    def prev_frame(self):
        if self.current_frame > 0:
            self.show_frame(self.current_frame-1)

    def choose_output_folder(self):
        folder_selected = filedialog.askdirectory(title="Select Output Folder")
        if folder_selected:
            self.export_to_excel(Path(folder_selected))
        else:
            messagebox.showerror("Folder Error", "No folder selected!")

    def export_to_excel(self, output_dir):
        first = self.data["Recipient Information"]["First Name"].get().strip()
        last = self.data["Recipient Information"]["Last Name"].get().strip()
        today = datetime.today()
        name_folder = f"{first}{last}_{today.strftime('%m%d%Y')}"
        folder_path = output_dir / name_folder
        folder_path.mkdir(parents=True, exist_ok=True)

        excel_path = folder_path / f"{last}_{today.strftime('%Y-%m-%d')}.xlsx"
        wb = Workbook(); ws = wb.active; ws.title = "M-SNAP Form Data"
        row = 1
        for section, fields in self.data.items():
            ws.cell(row=row, column=1, value=section)
            row += 1
            for key, var in fields.items():
                ws.cell(row=row, column=1, value=key)
                ws.cell(row=row, column=2, value=var.get())
                row += 1
            row += 1
        self.generate_documents(folder_path, first, last)
        self.append_to_voucher_tracker()
        messagebox.showinfo("Success", f"Saved everything to {folder_path}")
        self.restart_app()

    def fill_placeholders_in_docx(self, tmpl_path, out_path, ph):
        doc = Document(tmpl_path)

        def replace_text_in_paragraph(paragraph):
            if not paragraph.runs:
                return
            full_text = ''.join(run.text for run in paragraph.runs)
            for key, val in ph.items():
                full_text = full_text.replace(key, val)
            for run in paragraph.runs:
                run.clear()
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph)

        doc.save(out_path)

    def generate_documents(self, output_dir, first_name, last_name):
        script_dir = Path(__file__).resolve().parent
        base_dir = script_dir.parent
        today_str = datetime.today().strftime('%m/%d/%Y')
        today_file = datetime.today().strftime('%m%d%Y')

        templates = {
            "VOUCHER": "2025 Voucher Master.docx",
            "COVER": "2025 Cover Letter Master.docx",
            "LABEL": "Mailing Labels Master.docx"
        }

        recip = self.data["Recipient Information"]
        pet_no = 1
        for section in ["Pet 1 Information", "Pet 2 Information", "Pet 3 Information"]:
            pet = self.data[section]
            if not pet["Name"].get().strip():
                continue

            placeholders = {
                "«Date»": today_str,
                "«EXPIRES»": pet["Expires"].get(),
                "«Voucher_ID»": pet["Voucher"].get(),
                "«Customer_first_name»": first_name,
                "«Customer_last_name»": last_name,
                "«As_of_October_2010_Phone_Number»": recip["Day phone(s)"].get(),
                "«Customer_street_NEVER_abbreviate_except»": recip["Street, Apt # OR PO Box"].get(),
                "«Customer_city_CM__City_of_Morgantown»": recip["City"].get(),
                "«Cust_zip_code»": recip["Zip"].get(),
                "«Pet_Name»": pet["Name"].get(),
                "«Dog__Cat»": pet["Species"].get(),
                "«Spay_Neuter»": pet["Gender"].get(),
                "«Breed»": pet["Breed"].get(),
                "«Mailto_First_Name»": first_name,
                "«Mailto_Last_Name»": last_name,
                "«Street»": recip["Street, Apt # OR PO Box"].get(),
                "«City»": recip["City"].get(),
                "«Zip»": recip["Zip"].get(),
                "«Next Record»": ""
            }

            for doc_type, fname in templates.items():
                tmpl = base_dir / fname
                if not tmpl.exists():
                    tmpl = script_dir / fname
                if not tmpl.exists():
                    messagebox.showerror("Template Missing", f"Could not find template:\n· {base_dir/fname}\nor\n· {script_dir/fname}")
                    return

                out_name = output_dir / f"{last_name}_{today_file}_{doc_type}_{pet_no}.docx"

                # For labels: replace all placeholder instances, not just the first set
                if doc_type == "LABEL":
                    self.fill_all_labels_in_docx(str(tmpl), str(out_name), placeholders)
                else:
                    self.fill_placeholders_in_docx(str(tmpl), str(out_name), placeholders)

            pet_no += 1

    def fill_all_labels_in_docx(self, tmpl_path, out_path, ph):
        doc = Document(tmpl_path)

        def replace_placeholders_in_paragraph(paragraph, placeholders):
            full_text = ''.join(run.text for run in paragraph.runs)
            updated_text = full_text
            for key, value in placeholders.items():
                updated_text = updated_text.replace(key, value)

            if full_text != updated_text:
                for run in paragraph.runs:
                    run.text = ''  # Clear old content
                if paragraph.runs:
                    paragraph.runs[0].text = updated_text
                else:
                    paragraph.add_run(updated_text)

        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, ph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, ph)

        doc.save(out_path)

    def append_to_voucher_tracker(self):
        def find_first_empty_row(ws):
            row = 1
            while True:
                if all(
                    (ws.cell(row=row, column=col).value is None or str(ws.cell(row=row, column=col).value).strip() == "")
                    for col in range(1, 16)
                ):
                    return row
                row += 1
                while (row == 32 or row == 33):
                    row += 1
        script_dir = Path(__file__).resolve().parent
        default_path = script_dir / "M-SNAP VOUCHER TRACKER Master 2025.xlsm"

        if not default_path.exists():
            messagebox.showwarning(
                "Voucher Tracker File Not Found.",
                "Please select Voucher Tracker File"
            )
            file_path = filedialog.askopenfilename(
                title="Select the Voucher Tracker file",
                filetypes=[("Excel Macro-Enabled Workbook", "*.xlsm")]
            )
            if not file_path:
                messagebox.showerror("Operation Cancelled", "No voucher tracker selected. Data will NOT be appended.")
                return
            path = Path(file_path)
        else:
            path = default_path

        try:
            wb = load_workbook(path, keep_vba=True)
        
            if "Vouchers" in wb.sheetnames:
                ws = wb["Vouchers"]
            else:
                ws = wb.active

            row = find_first_empty_row(ws)

            recip = self.data["Recipient Information"]

            for section_name in ["Pet 1 Information", "Pet 2 Information", "Pet 3 Information"]:
                pet = self.data[section_name]
                if not pet["Name"].get().strip():
                    continue

                ws.cell(row=row, column=1, value=datetime.today())
                ws.cell(row=row, column=2, value=pet["Voucher"].get())  # Voucher ID
                ws.cell(row=row, column=3, value=recip["First Name"].get())
                ws.cell(row=row, column=4, value=recip["Last Name"].get())
                ws.cell(row=row, column=5, value=pet["Name"].get())     # Pet Name
                ws.cell(row=row, column=6, value=recip["City"].get())
                ws.cell(row=row, column=7, value=recip["Zip"].get())
                ws.cell(row=row, column=8, value=pet["Species"].get())  # Species
                gender = pet["Gender"].get()
                gender_export = "N" if gender == "Male" else ("S" if gender == "Female" else "")
                ws.cell(row=row, column=9, value=gender_export)
                ws.cell(row=row, column=10, value=pet["Stray?"].get())   
                ws.cell(row=row, column=14, value=recip["Day phone(s)"].get())
                ws.cell(row=row, column=15, value=recip["How did you hear about M-SNAP?"].get())
                row += 1

            wb.save(path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to append to the voucher tracker:\n{str(e)}")

    def restart_app(self):
        for sec, fields in self.data.items():
            for key, var in fields.items():
                if sec.startswith("Pet") and key in self.dropdown_fields_common:
                    var.set("N/A")
                elif sec.startswith("Pet") and key in self.gender_field:
                    var.set("N/A")
                elif sec == "Recipient Information" and key in self.yes_no_field:
                    var.set("No")
                else:
                    var.set("")
        for flag in self.add_another_pet_flags.values():
            flag.set(0)
        self.show_frame(0)

if __name__ == "__main__":
    app = MSnapApp()
    app.mainloop()
